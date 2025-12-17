#!/usr/bin/env python3
"""
Color-Code Geniza Letters in Word Document

This script reads JSON codings files, extracts original letter transcriptions,
maps categories to character positions, and generates a Word document with
color-coded text based on categories.

Usage:
    python color_code_letters.py --json ib_geniza_codings_sandbox_gpt4o_mini.json [options]
"""

import argparse
import json
import os
import re
import sys
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH


def read_json_codings(json_path: str, model_name: Optional[str] = None) -> Dict:
    """
    Read JSON codings file and extract segments for each document.

    Args:
        json_path: Path to JSON codings file
        model_name: Optional model name to filter (if multiple models in JSON)

    Returns:
        Dictionary mapping document_id to segments data
    """
    print(f"Reading JSON file: {json_path}")
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Extract segments for each document
    documents_data: Dict = {}
    for doc_id_str, doc_data in data.items():
        doc_id = int(doc_id_str)
        segments_data = doc_data.get("segments", {})

        # Handle different JSON structures
        if "base_segments" in segments_data:
            segments = segments_data["base_segments"]
        elif isinstance(segments_data, list):
            segments = segments_data
        else:
            # Try to find segments in model_codings
            if "model_codings" in segments_data and model_name:
                segments = segments_data["model_codings"].get(model_name, [])
            else:
                segments = []

        # Extract codings (percentages) - handle different structures
        codings = {}
        if "codings" in doc_data:
            codings_data = doc_data["codings"]
            # If codings is a dict with model names, get first model or specified model
            if isinstance(codings_data, dict):
                if model_name and model_name in codings_data:
                    codings = codings_data[model_name]
                elif codings_data:
                    # Get first model's codings
                    codings = list(codings_data.values())[0]
            else:
                codings = codings_data

        documents_data[doc_id] = {
            "segments": segments,
            "total_characters": doc_data.get("total_characters", 0),
            "mode": doc_data.get("mode", "single"),
            "codings": codings,
        }

    print(f"Found {len(documents_data)} documents in JSON")
    return documents_data


def extract_letters_from_csv(csv_path: str, document_ids: List[int]) -> Dict[int, str]:
    """
    Extract Judeo-Arabic content from CSV for given document IDs.

    Args:
        csv_path: Path to the CSV file
        document_ids: List of document IDs to extract

    Returns:
        Dictionary mapping document_id to full letter content (Judeo-Arabic text)
    """
    print(f"Reading CSV file: {csv_path}")
    df = pd.read_csv(csv_path)

    # Filter by document_ids
    df_filtered = df[df["document_id"].isin(document_ids)].copy()

    if df_filtered.empty:
        print("Warning: No documents found for the provided document IDs")
        return {}

    # Group by document_id and combine content
    letters: Dict[int, str] = {}
    for doc_id in document_ids:
        doc_rows = df_filtered[df_filtered["document_id"] == doc_id]

        # Extract content from rows that have content
        content_parts: List[str] = []
        for _, row in doc_rows.iterrows():
            content = str(row.get("content", "")).strip()
            if content and content != "nan":
                # Check if content contains Hebrew script (Judeo-Arabic)
                if re.search(r"[\u0590-\u05FF]", content):
                    content_parts.append(content)

        if content_parts:
            # Join content parts with newlines
            full_content = "\n".join(content_parts)
            letters[doc_id] = full_content
            print(f"Extracted {len(full_content)} characters for document_id {doc_id}")
        else:
            print(f"Warning: No Judeo-Arabic content found for document_id {doc_id}")

    return letters


def map_categories_to_positions(
    letter_text: str, segments: List[Dict], default_category: str = "Uncategorized"
) -> List[Optional[str]]:
    """
    Map each character position in the letter to a category based on segments.

    Args:
        letter_text: Full letter text
        segments: List of segment dictionaries with start_char, end_char, category
        default_category: Category to use for unassigned characters

    Returns:
        List where each index corresponds to a character position and value is category name
    """
    category_map: List[Optional[str]] = [None] * len(letter_text)

    # Process segments in order
    for segment in segments:
        start_char = segment.get("start_char", 0)
        end_char = segment.get("end_char", start_char)
        category = segment.get("category")

        # Skip if category is None/null
        if category is None:
            category = default_category

        # Ensure indices are within bounds
        start_char = max(0, min(start_char, len(letter_text)))
        end_char = max(start_char, min(end_char, len(letter_text)))

        # Assign category to character positions (first segment wins for overlaps)
        for i in range(start_char, end_char):
            if category_map[i] is None:
                category_map[i] = category

    # Fill in any remaining None values with default category
    for i in range(len(category_map)):
        if category_map[i] is None:
            category_map[i] = default_category

    return category_map


def assign_colors_to_categories(categories: List[str]) -> Dict[str, RGBColor]:
    """
    Assign distinct colors to each category.

    Args:
        categories: List of unique category names

    Returns:
        Dictionary mapping category name to RGBColor
    """
    # Predefined palette of 10+ distinct colors
    color_palette = [
        RGBColor(31, 119, 180),   # Blue
        RGBColor(255, 127, 14),   # Orange
        RGBColor(44, 160, 44),    # Green
        RGBColor(214, 39, 40),    # Red
        RGBColor(148, 103, 189),  # Purple
        RGBColor(140, 86, 75),    # Brown
        RGBColor(227, 119, 194),  # Pink
        RGBColor(127, 127, 127),  # Gray
        RGBColor(188, 189, 34),   # Olive
        RGBColor(23, 190, 207),   # Cyan
        RGBColor(255, 187, 120),  # Light Orange
        RGBColor(152, 223, 138),  # Light Green
        RGBColor(174, 199, 232),  # Light Blue
        RGBColor(255, 152, 150),  # Light Red
        RGBColor(197, 176, 213),  # Light Purple
    ]

    # Default color for uncategorized
    default_color = RGBColor(0, 0, 0)  # Black

    category_colors: Dict[str, RGBColor] = {}
    color_index = 0

    for category in categories:
        if category == "Uncategorized":
            category_colors[category] = default_color
        else:
            category_colors[category] = color_palette[color_index % len(color_palette)]
            color_index += 1

    return category_colors


def generate_word_document(
    documents_data: Dict,
    letters: Dict[int, str],
    output_path: str,
    category_colors: Dict[str, RGBColor],
) -> None:
    """
    Generate Word document with color-coded letters.

    Args:
        documents_data: Dictionary mapping document_id to segments data
        letters: Dictionary mapping document_id to letter text
        output_path: Path to output Word document
        category_colors: Dictionary mapping category name to RGBColor
    """
    print(f"\nGenerating Word document: {output_path}")
    doc = Document()

    # Add title
    title = doc.add_heading("Color-Coded Geniza Letters", 0)
    title.alignment = 1  # Center alignment

    # Add legend
    doc.add_heading("Category Color Legend", level=1)
    legend_table = doc.add_table(rows=1, cols=2)
    legend_table.style = "Light Grid Accent 1"
    header_cells = legend_table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Color"

    # Sort categories for consistent legend
    sorted_categories = sorted(category_colors.keys())
    for category in sorted_categories:
        row_cells = legend_table.add_row().cells
        row_cells[0].text = category
        # Add colored cell
        color_cell = row_cells[1]
        color_cell.text = "█" * 10  # Color block
        for paragraph in color_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = category_colors[category]

    doc.add_page_break()

    # Process each letter
    for doc_id in sorted(letters.keys()):
        if doc_id not in documents_data:
            print(f"Warning: No segments found for document_id {doc_id}, skipping...")
            continue

        letter_text = letters[doc_id]
        segments = documents_data[doc_id]["segments"]
        codings = documents_data[doc_id].get("codings", {})

        # Add document ID heading
        doc.add_heading(f"Document ID: {doc_id}", level=1)

        # Add category percentages table if available
        if codings:
            doc.add_heading("Category Percentages", level=2)
            percentages_table = doc.add_table(rows=1, cols=2)
            percentages_table.style = "Light Grid Accent 1"
            header_cells = percentages_table.rows[0].cells
            header_cells[0].text = "Category"
            header_cells[1].text = "Percentage"

            # Sort categories by percentage (descending) for better readability
            sorted_codings = sorted(
                codings.items(), key=lambda x: x[1], reverse=True
            )
            for category, percentage in sorted_codings:
                row_cells = percentages_table.add_row().cells
                row_cells[0].text = category
                row_cells[1].text = f"{percentage * 100:.2f}%"
                # Color the category name cell with the category color
                if category in category_colors:
                    for paragraph in row_cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = category_colors[category]

            doc.add_paragraph()  # Add spacing

        # Map categories to character positions
        category_map = map_categories_to_positions(letter_text, segments)

        # Split text by newlines to preserve paragraph structure
        paragraphs_text = letter_text.split("\n")
        char_index = 0

        for para_text in paragraphs_text:
            if not para_text:
                # Empty line - add empty paragraph (right-aligned)
                empty_para = doc.add_paragraph()
                empty_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                char_index += 1  # Account for the newline character
                continue

            # Create a paragraph for this line and right-align it
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Process characters in this paragraph, grouping by category
            current_category = None
            current_text = ""

            for char in para_text:
                category = category_map[char_index]

                if category != current_category:
                    # Category changed: add current run and start new one
                    if current_text:
                        run = paragraph.add_run(current_text)
                        if current_category and current_category in category_colors:
                            run.font.color.rgb = category_colors[current_category]
                        current_text = ""

                    current_category = category
                    current_text = char
                else:
                    # Same category: append to current text
                    current_text += char

                char_index += 1

            # Add final run for this paragraph if any text remains
            if current_text:
                run = paragraph.add_run(current_text)
                if current_category and current_category in category_colors:
                    run.font.color.rgb = category_colors[current_category]

            # Account for the newline character
            char_index += 1

        # Add page break between letters
        doc.add_page_break()

    # Save document
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word document with color-coded Geniza letters"
    )
    parser.add_argument(
        "--json",
        default="ib_geniza_codings_sandbox_gpt4o_mini.json",
        help="Path to JSON codings file (default: ib_geniza_codings_sandbox_gpt4o_mini.json)",
    )
    parser.add_argument(
        "--csv",
        default="princetongenizalab_pgp-metadata__data_footnotes-csv__11_29_2025.csv",
        help=(
            "Path to CSV file "
            "(default: princetongenizalab_pgp-metadata__data_footnotes-csv__11_29_2025.csv)"
        ),
    )
    parser.add_argument(
        "--output",
        default="color_coded_letters.docx",
        help="Output Word document path (default: color_coded_letters.docx)",
    )
    parser.add_argument(
        "--model",
        default=None,
        help="Model name to use if multiple models in JSON (default: first available)",
    )

    args = parser.parse_args()

    # Read JSON codings
    documents_data = read_json_codings(args.json, args.model)

    if not documents_data:
        print("Error: No documents found in JSON file")
        sys.exit(1)

    # Extract document IDs
    document_ids = list(documents_data.keys())
    print(f"Processing {len(document_ids)} documents")

    # Extract original letters from CSV
    letters = extract_letters_from_csv(args.csv, document_ids)

    if not letters:
        print("Error: No letters extracted from CSV")
        sys.exit(1)

    # Collect all unique categories
    all_categories = set()
    for doc_id, data in documents_data.items():
        for segment in data["segments"]:
            category = segment.get("category")
            if category:
                all_categories.add(category)
    all_categories.add("Uncategorized")  # Add default category

    print(f"Found {len(all_categories)} unique categories")

    # Assign colors to categories
    category_colors = assign_colors_to_categories(sorted(all_categories))

    print("\nCategory color assignments:")
    for category in sorted(category_colors.keys()):
        print(f"  {category}")

    # Generate Word document
    generate_word_document(documents_data, letters, args.output, category_colors)

    print("\nDone!")


if __name__ == "__main__":
    main()
